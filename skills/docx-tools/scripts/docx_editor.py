#!/usr/bin/env python3
"""
编辑 Word 文档
"""

from docx import Document
import sys
import json

def edit_document(file_path, edits):
    """编辑 Word 文档"""

    try:
        doc = Document(file_path)

        # 替换文本
        if 'replacements' in edits:
            for para in doc.paragraphs:
                for old, new in edits['replacements'].items():
                    if old in para.text:
                        para.text = para.text.replace(old, new)

        # 在末尾添加内容
        if 'append_text' in edits:
            doc.add_paragraph(edits['append_text'])

        # 添加标题
        if 'add_heading' in edits:
            level = edits['add_heading'].get('level', 1)
            doc.add_heading(edits['add_heading']['text'], level)

        output_path = edits.get('output_path', file_path)
        doc.save(output_path)
        return {'success': True, 'output_path': output_path}

    except Exception as e:
        return {'error': str(e)}

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(json.dumps({
            'error': 'Usage: docx_editor.py <file_path> <json_edits> [output_path]'
        }, ensure_ascii=False))
        sys.exit(1)

    file_path = sys.argv[1]
    json_edits = sys.argv[2]
    output_path = sys.argv[3] if len(sys.argv) > 3 else file_path

    try:
        edits = json.loads(json_edits)
        edits['output_path'] = output_path
        result = edit_document(file_path, edits)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    except json.JSONDecodeError as e:
        print(json.dumps({'error': f'Invalid JSON: {str(e)}'}, ensure_ascii=False))

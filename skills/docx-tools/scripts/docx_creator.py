#!/usr/bin/env python3
"""
创建 Word 文档
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import json

def create_document(content):
    """创建 Word 文档"""

    try:
        doc = Document()

        # 标题
        if 'title' in content:
            doc.add_heading(content['title'], 0)

        # 段落
        if 'paragraphs' in content:
            for para in content['paragraphs']:
                p = doc.add_paragraph(para['text'])
                if 'style' in para:
                    p.style = para['style']

        # 表格
        if 'tables' in content:
            for table_data in content['tables']:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_data):
                    for j, cell_data in enumerate(row_data):
                        table.rows[i].cells[j].text = cell_data

        output_path = content.get('output_path', '/tmp/output.docx')
        doc.save(output_path)
        return {'success': True, 'output_path': output_path}

    except Exception as e:
        return {'error': str(e)}

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'Usage: docx_creator.py <json_content> [output_path]'}, ensure_ascii=False))
        sys.exit(1)

    json_content = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else '/tmp/output.docx'

    try:
        content = json.loads(json_content)
        content['output_path'] = output_path
        result = create_document(content)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    except json.JSONDecodeError as e:
        print(json.dumps({'error': f'Invalid JSON: {str(e)}'}, ensure_ascii=False))

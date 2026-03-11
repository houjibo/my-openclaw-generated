#!/usr/bin/env python3
"""
读取 Word 文档内容
"""

from docx import Document
import sys
import json

def read_docx(file_path):
    """读取 Word 文档"""
    try:
        doc = Document(file_path)

        result = {
            'paragraphs': [],
            'tables': [],
            'metadata': {}
        }

        # 读取段落
        for para in doc.paragraphs:
            if para.text.strip():
                result['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                })

        # 读取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            result['tables'].append(table_data)

        # 元数据
        result['metadata'] = {
            'paragraphs_count': len(doc.paragraphs),
            'tables_count': len(doc.tables)
        }

        return result

    except Exception as e:
        return {'error': str(e)}

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(json.dumps({'error': 'Usage: docx_reader.py <file_path>'}, ensure_ascii=False))
        sys.exit(1)

    file_path = sys.argv[1]
    result = read_docx(file_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
